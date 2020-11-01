# -*- coding: utf-8 -*-
"""
Created on Sun Oct 30 10:21:16 2020
@author: xq

preparation
1. install python
   https://www.python.org/downloads/release/python-390/
   Windows x86-64 executable installer
2. install python-docx
   1) run cmd
   2) in CMD run: pip install python-docx
3. pip install matplotlib
   1) run cmd
   2) in CMD run: pip install matplotlib

usage
   1) run cmd
   2) change the file name in line27 to the rinex obs file
   3) in cmd run: python cpt3.py
"""
from math import nan, isnan
from matplotlib import pyplot as plt 
import docx
import os.path

full_file_name = r'data\data.obs_'

CURRENT_HEADER = []
DIC_SET = {}
STAS_NOT_CHANGED_HDR_FIELD_COUNT = 2
SKIP_ROW_COUNT = 2
SAT_OFFSET = 2

def init_database() -> None:
    for inx in range(40):
        DIC_SET['C{0:02d}'.format(inx+1)]=[]
        DIC_SET['G{0:02d}'.format(inx+1)]=[]

def read_cpt3(filename: str) -> None:
    with open(filename) as f:
        for inx, line in enumerate(f.readlines()):
            parse_cpt3(inx, line)

def parse_cpt3(inx: int, line: str) -> None:
    if inx<SKIP_ROW_COUNT:
        return
    if inx%2==0:
        parse_header(line)
    else:
        parse_body(line)

def parse_header(line: str) -> None:
    global CURRENT_HEADER
    hdr = line.split()
    if len(hdr)==STAS_NOT_CHANGED_HDR_FIELD_COUNT:
        CURRENT_HEADER[0] = hdr[0]
    else:
        CURRENT_HEADER = hdr

def parse_body(line: str) -> None:
    for inx, val in enumerate(line.split()):
        try:
            DIC_SET[CURRENT_HEADER[inx+SAT_OFFSET]].append(float(val))
        except:
            DIC_SET[CURRENT_HEADER[inx+SAT_OFFSET]].append(0)
        else:
            DIC_SET[CURRENT_HEADER[inx+SAT_OFFSET]].append(0)

    for k in DIC_SET.keys():
        if k not in CURRENT_HEADER:
            DIC_SET[k].append(nan)

def all_nan(ls) -> bool:
    for val in ls:
        if not isnan(val):
            return False
    return True

def plot_data(file_name: str) -> None:
    for k in DIC_SET.keys():
        if all_nan(DIC_SET[k]):
            continue
        plt.plot(DIC_SET[k],label=k)
    plt.legend(loc='upper center',ncol=6)
    plt.title(file_name)
    plt.grid(True)
    plt.savefig("a.png")
    plt.clf()

def savefig_to_report(doc, file_name: str) -> None:
    if os.path.isfile(file_name):
        pass
    else:
        doc.add_paragraph("could not get data to plot.")
        return

    init_database()
    read_cpt3(file_name)
    plot_data(file_name)
    doc.add_picture('a.png')

if __name__ == '__main__':
    file_name = os.path.splitext(full_file_name)[0]
    os.system(r'teqc.exe +qc +plot {} > out.S'.format(full_file_name))

    doc = docx.Document()
    doc.add_heading(r"TEQC Base GNSS Data Quality Report", level=1)
    doc.add_heading(r"Summary",level=2)
    with open(r'out.S') as f:
        for line in f.readlines():
            line = line.strip() 
            doc.add_paragraph(line)

    doc.add_heading(r"Plot",   level=2)

    doc.add_paragraph("Ionosphere delay rate" , style='ListNumber')
    savefig_to_report(doc, r'{}.d12'.format(file_name))
    doc.add_paragraph("Ionosphere delay" , style='ListNumber')
    savefig_to_report(doc, r'{}.i12'.format(file_name))
    doc.add_paragraph("multipath" , style='ListNumber')
    savefig_to_report(doc, r'{}.m12'.format(file_name))
    doc.add_paragraph("multipath" , style='ListNumber')
    savefig_to_report(doc, r'{}.m21'.format(file_name))
    doc.add_paragraph("azi" , style='ListNumber')
    savefig_to_report(doc, r'{}.azi'.format(file_name))
    doc.add_paragraph("ele" , style='ListNumber')
    savefig_to_report(doc, r'{}.ele'.format(file_name))
    doc.add_paragraph("sn1" , style='ListNumber')
    savefig_to_report(doc, r'{}.sn1'.format(file_name))
    doc.add_paragraph("sn2" , style='ListNumber')
    savefig_to_report(doc, r'{}.sn2'.format(file_name))

    doc.save(r'{}.docx'.format(file_name))


